from __future__ import annotations

import math
import os
import tempfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "APPLICATION_ARCHITECTURE_AND_CLASS_FLOW.docx"
ASSET_DIR = Path(tempfile.gettempdir()) / "fie_architecture_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

NAVY = "17324D"
BLUE = "2E74B5"
TEAL = "16827C"
GREEN = "3F7D55"
GOLD = "A66A00"
RED = "A23B3B"
INK = "1F2933"
MUTED = "5D6B78"
PALE_BLUE = "EAF2F8"
PALE_TEAL = "E8F5F3"
PALE_GREEN = "EDF6EF"
PALE_GOLD = "FFF5E3"
PALE_RED = "FBEDEE"
LIGHT = "F4F6F8"
WHITE = "FFFFFF"
BORDER = "CAD3DC"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 100, "bottom": 100, "start": 120, "end": 120}


def rgb(value: str) -> RGBColor:
    value = value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=None):
    margins = margins or CELL_MARGIN_DXA
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[edge]))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=False):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("FINANCIAL INTELLIGENCE ENGINE  |  ARCHITECTURE GUIDE")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_paragraph_keep(p, keep_next=True)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first, rest = text.split(":", 1)
        r1 = p.add_run(first + ":")
        set_run_font(r1, bold=True)
        r2 = p.add_run(rest)
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED, italic=True)
    return p


def add_diagram(doc, path, alt_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(path), width=Inches(6.45))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)
    return p


def add_callout(doc, title, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    set_table_borders(table, color=accent, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=font_size, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    return table


def page_break(doc):
    doc.add_page_break()


def font_path(bold=False):
    candidates = [
        Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts" / ("arialbd.ttf" if bold else "arial.ttf"),
        Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts" / ("calibrib.ttf" if bold else "calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def pil_font(size, bold=False):
    path = font_path(bold)
    if path:
        return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def wrap_lines(draw, text, font, max_width):
    words = str(text).split()
    lines = []
    current = ""
    for word in words:
        trial = word if not current else current + " " + word
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw, box, title, body=None, fill=PALE_BLUE, outline=BLUE, title_size=29, body_size=22):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=14, fill="#" + fill, outline="#" + outline, width=4)
    title_font = pil_font(title_size, bold=True)
    body_font = pil_font(body_size)
    title_lines = wrap_lines(draw, title, title_font, x2 - x1 - 30)
    y = y1 + 16
    for line in title_lines:
        bbox = draw.textbbox((0, 0), line, font=title_font)
        draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, font=title_font, fill="#" + NAVY)
        y += title_size + 5
    if body:
        y += 4
        for line in wrap_lines(draw, body, body_font, x2 - x1 - 32):
            bbox = draw.textbbox((0, 0), line, font=body_font)
            draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, font=body_font, fill="#" + INK)
            y += body_size + 5


def draw_label(draw, xy, text, size=24, color=MUTED, bold=False, anchor="la"):
    draw.text(xy, text, font=pil_font(size, bold=bold), fill="#" + color, anchor=anchor)


def draw_arrow(draw, start, end, color=BLUE, width=5):
    draw.line([start, end], fill="#" + color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    spread = math.pi / 7
    p1 = (
        end[0] - length * math.cos(angle - spread),
        end[1] - length * math.sin(angle - spread),
    )
    p2 = (
        end[0] - length * math.cos(angle + spread),
        end[1] - length * math.sin(angle + spread),
    )
    draw.polygon([end, p1, p2], fill="#" + color)


def save_diagram(image, name):
    path = ASSET_DIR / name
    image.save(path, format="PNG", dpi=(180, 180))
    return path


def create_system_architecture():
    image = Image.new("RGB", (1800, 1130), "white")
    draw = ImageDraw.Draw(image)
    draw_label(draw, (900, 42), "Phase 1 System Architecture", 42, NAVY, True, "ma")

    draw_box(draw, (620, 90, 1180, 185), "User Browser", "Dashboard views and investment question", PALE_BLUE, BLUE)
    draw_arrow(draw, (900, 185), (900, 225))
    draw_box(draw, (570, 225, 1230, 335), "Frontend and Local API", "HTML / CSS / JavaScript + DashboardHandler on port 8000", LIGHT, MUTED)
    draw_arrow(draw, (900, 335), (900, 380))
    draw_box(draw, (620, 380, 1180, 485), "OrchestratorAgent", "Ticker selection, sequencing, result aggregation", PALE_TEAL, TEAL)

    agent_boxes = [
        ((75, 560, 405, 690), "RetrieverAgent", "ChromaDB first; hybrid local fallback", PALE_BLUE, BLUE),
        ((485, 560, 815, 690), "SentimentAgent", "News + retrieved evidence sentiment", PALE_GOLD, GOLD),
        ((985, 560, 1315, 690), "RiskAgent", "Model-backed risk classification", PALE_RED, RED),
        ((1395, 560, 1725, 690), "ForecastAgent", "30-day direction classification", PALE_GREEN, GREEN),
    ]
    for box, title, body, fill, outline in agent_boxes:
        draw_box(draw, box, title, body, fill, outline, title_size=27, body_size=20)
        draw_arrow(draw, (900, 485), ((box[0] + box[2]) // 2, box[1]), color=outline, width=4)

    draw_box(draw, (570, 775, 1230, 875), "DecisionAgent", "30% sentiment + 35% inverse risk + 35% forecast return", PALE_GOLD, GOLD)
    for box, _, _, _, outline in agent_boxes:
        draw_arrow(draw, ((box[0] + box[2]) // 2, box[3]), (900, 775), color=outline, width=4)
    draw_arrow(draw, (900, 875), (900, 910), color=TEAL)
    draw_box(draw, (570, 910, 1230, 1005), "ExplainabilityAgent + GenAIProvider", "Grounded report, evidence, limitations, disclaimer", PALE_TEAL, TEAL)
    draw_arrow(draw, (900, 1005), (900, 1060), color=BLUE)
    draw_label(draw, (900, 1090), "Final JSON response, dashboard presentation, and Markdown report", 27, NAVY, True, "ma")

    draw_label(draw, (75, 760), "Persistent intelligence assets", 22, MUTED, True)
    draw_label(draw, (75, 795), "Chroma vectors | CSV features | Pickle models | Reports", 20, MUTED)
    return save_diagram(image, "01_system_architecture.png")


def create_runtime_flow():
    image = Image.new("RGB", (1800, 1120), "white")
    draw = ImageDraw.Draw(image)
    draw_label(draw, (900, 40), "Runtime Query and Decision Flow", 42, NAVY, True, "ma")

    stages = [
        ((70, 110, 430, 215), "1. Browser", "POST /api/ask", PALE_BLUE, BLUE),
        ((520, 110, 880, 215), "2. DashboardHandler", "Validate input; extract tickers", LIGHT, MUTED),
        ((970, 110, 1330, 215), "3. OrchestratorAgent", "Create one coordinated run", PALE_TEAL, TEAL),
        ((1420, 110, 1760, 215), "4. RetrieverAgent", "Top-k grounded evidence", PALE_BLUE, BLUE),
    ]
    for i, item in enumerate(stages):
        draw_box(draw, *item)
        if i < len(stages) - 1:
            draw_arrow(draw, (item[0][2], 162), (stages[i + 1][0][0], 162))

    draw_arrow(draw, (1590, 215), (1590, 300), color=BLUE)
    draw_box(draw, (1340, 300, 1760, 415), "5. SentimentAgent", "Evidence + demo news -> score, label, trend", PALE_GOLD, GOLD)
    draw_arrow(draw, (1340, 357), (1030, 357), color=GOLD)

    draw_box(draw, (610, 300, 1030, 415), "6A. RiskAgent", "Load risk model and ticker feature row", PALE_RED, RED)
    draw_box(draw, (610, 500, 1030, 615), "6B. ForecastAgent", "Load forecast model and ticker feature row", PALE_GREEN, GREEN)
    draw_arrow(draw, (820, 415), (820, 500), color=TEAL)
    draw_label(draw, (1070, 525), "Forecast is logically independent;\ncurrent orchestrator executes calls sequentially.", 21, MUTED)

    draw_box(draw, (610, 700, 1030, 815), "7. DecisionAgent", "Weighted score -> Strong Buy / Buy / Hold / Sell / Strong Sell", PALE_GOLD, GOLD)
    draw_arrow(draw, (820, 615), (820, 700), color=GREEN)
    draw.line([(610, 357), (520, 357), (520, 757)], fill="#" + RED, width=5)
    draw_arrow(draw, (520, 757), (610, 757), color=RED)
    draw_label(draw, (250, 525), "Risk and forecast outputs\nplus sentiment signals", 21, MUTED, True, "ma")

    draw_box(draw, (610, 885, 1030, 1000), "8. ExplainabilityAgent", "Assemble evidence-grounded report; optional OpenAI synthesis", PALE_TEAL, TEAL)
    draw_arrow(draw, (820, 815), (820, 885), color=TEAL)
    draw_arrow(draw, (1030, 942), (1390, 942), color=BLUE)
    draw_box(draw, (1390, 885, 1760, 1000), "9. API Response", "Report + agent summaries + suggestions", PALE_BLUE, BLUE)
    draw_arrow(draw, (1390, 1000), (430, 1050), color=MUTED, width=3)
    draw_label(draw, (250, 1050), "Rendered in browser", 24, NAVY, True, "ma")
    return save_diagram(image, "02_runtime_flow.png")


def create_rag_flow():
    image = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(image)
    draw_label(draw, (900, 42), "Persistent RAG Lifecycle", 42, NAVY, True, "ma")
    draw_label(draw, (350, 100), "INDEXING PATH", 25, BLUE, True, "ma")
    draw_label(draw, (1450, 100), "QUERY PATH", 25, TEAL, True, "ma")

    left = [
        ((80, 150, 620, 255), "Documents", "PDF, MD, TXT, CSV, JSON", PALE_BLUE, BLUE),
        ((80, 315, 620, 420), "DocumentLoader", "Extract text; collect supported files", LIGHT, MUTED),
        ((80, 480, 620, 585), "Chunking", "900 characters with 120-character overlap", PALE_GOLD, GOLD),
        ((80, 645, 620, 750), "HashingEmbedder", "Deterministic normalized 384-dimensional vectors", PALE_TEAL, TEAL),
    ]
    for i, item in enumerate(left):
        draw_box(draw, *item)
        if i < len(left) - 1:
            draw_arrow(draw, (350, item[0][3]), (350, left[i + 1][0][1]))

    draw_box(draw, (690, 395, 1110, 610), "ChromaDB", "PersistentClient\ncollection: phase1_financial_docs\ndata/vectors/chroma/", PALE_GREEN, GREEN, title_size=34, body_size=24)
    draw_arrow(draw, (620, 697), (690, 560), color=GREEN)

    right = [
        ((1180, 150, 1720, 255), "User Question", "Natural-language financial or project query", PALE_BLUE, BLUE),
        ((1180, 315, 1720, 420), "Query Embedding", "Same HashingEmbedder contract", PALE_TEAL, TEAL),
        ((1180, 645, 1720, 750), "Top-k Evidence", "Text, source metadata, score, distance", PALE_GOLD, GOLD),
        ((1180, 810, 1720, 930), "Answer Layer", "Local extractive answer or OpenAI grounded synthesis", PALE_BLUE, BLUE),
    ]
    for i, item in enumerate(right):
        draw_box(draw, *item)
    draw_arrow(draw, (1450, 255), (1450, 315))
    draw_arrow(draw, (1180, 367), (1110, 455), color=TEAL)
    draw_arrow(draw, (1110, 550), (1180, 697), color=GREEN)
    draw_arrow(draw, (1450, 750), (1450, 810), color=BLUE)
    draw_label(draw, (900, 980), "RetrieverAgent uses ChromaDB when indexed; local TF-IDF/keyword retrieval remains a resilience fallback.", 23, MUTED, True, "ma")
    return save_diagram(image, "03_rag_flow.png")


def create_ml_flow():
    image = Image.new("RGB", (1800, 1090), "white")
    draw = ImageDraw.Draw(image)
    draw_label(draw, (900, 40), "Machine Learning: Training and Inference", 42, NAVY, True, "ma")
    draw_label(draw, (50, 100), "OFFLINE BUILD / TRAINING", 25, BLUE, True)
    draw.rounded_rectangle((40, 135, 1760, 560), radius=18, outline="#" + BORDER, width=3, fill="#FAFBFC")

    boxes = [
        ((75, 200, 365, 335), "US Equity Universe Builder", "Nasdaq + SEC cached universe", PALE_BLUE, BLUE),
        ((430, 200, 720, 335), "Phase1DataPipeline", "10,000 x 20 feature/label table", PALE_TEAL, TEAL),
        ((785, 200, 1075, 335), "Phase1ModelTrainer", "80/20 split and evaluation", PALE_GOLD, GOLD),
        ((1140, 155, 1450, 285), "RiskScorer", "Low / Medium / High", PALE_RED, RED),
        ((1140, 360, 1450, 490), "PricePredictor", "Down / Flat / Up", PALE_GREEN, GREEN),
        ((1510, 200, 1725, 445), "Artifacts", "PKL models\nJSON metrics", LIGHT, MUTED),
    ]
    for item in boxes:
        draw_box(draw, *item, title_size=25, body_size=20)
    draw_arrow(draw, (365, 267), (430, 267))
    draw_arrow(draw, (720, 267), (785, 267))
    draw_arrow(draw, (1075, 267), (1140, 220), color=RED)
    draw_arrow(draw, (1075, 267), (1140, 425), color=GREEN)
    draw_arrow(draw, (1450, 220), (1510, 270), color=RED)
    draw_arrow(draw, (1450, 425), (1510, 380), color=GREEN)

    draw_label(draw, (50, 615), "ONLINE INFERENCE", 25, TEAL, True)
    draw.rounded_rectangle((40, 650, 1760, 1010), radius=18, outline="#" + BORDER, width=3, fill="#FAFBFC")
    infer = [
        ((75, 755, 380, 895), "Ticker Feature Row", "Read from feature CSV", PALE_BLUE, BLUE),
        ((465, 700, 800, 825), "RiskAgent", "RiskScorer.predict_proba()", PALE_RED, RED),
        ((465, 850, 800, 975), "ForecastAgent", "PricePredictor.predict_proba()", PALE_GREEN, GREEN),
        ((900, 755, 1215, 895), "DecisionAgent", "Combine model signals", PALE_GOLD, GOLD),
        ((1315, 755, 1715, 895), "Explainable Output", "Label, probabilities, score, recommendation", PALE_TEAL, TEAL),
    ]
    for item in infer:
        draw_box(draw, *item, title_size=27, body_size=20)
    draw_arrow(draw, (380, 825), (465, 762), color=RED)
    draw_arrow(draw, (380, 825), (465, 912), color=GREEN)
    draw_arrow(draw, (800, 762), (900, 825), color=RED)
    draw_arrow(draw, (800, 912), (900, 825), color=GREEN)
    draw_arrow(draw, (1215, 825), (1315, 825), color=TEAL)
    draw_label(draw, (900, 1050), "Current labels are learned from deterministic Phase 1 proxy features; production training requires time-stamped real market outcomes.", 22, RED, True, "ma")
    return save_diagram(image, "04_ml_flow.png")


def create_class_flow():
    image = Image.new("RGB", (1800, 1200), "white")
    draw = ImageDraw.Draw(image)
    draw_label(draw, (900, 38), "Runtime Class Dependency Flow", 42, NAVY, True, "ma")

    draw_box(draw, (600, 90, 1200, 185), "DashboardHandler", "HTTP routes, static assets, JSON response", LIGHT, MUTED)
    draw_arrow(draw, (900, 185), (900, 235))
    draw_box(draw, (600, 235, 1200, 335), "OrchestratorAgent", "Owns and coordinates specialist agent instances", PALE_TEAL, TEAL)

    agents = [
        ((45, 420, 325, 535), "RetrieverAgent", "returns AgentResult + EvidenceItem", PALE_BLUE, BLUE),
        ((390, 420, 670, 535), "SentimentAgent", "returns ticker sentiment", PALE_GOLD, GOLD),
        ((735, 420, 1015, 535), "RiskAgent", "returns risk prediction", PALE_RED, RED),
        ((1080, 420, 1360, 535), "ForecastAgent", "returns direction prediction", PALE_GREEN, GREEN),
        ((1425, 420, 1755, 535), "DecisionAgent", "returns investment decision", PALE_GOLD, GOLD),
    ]
    for item in agents:
        draw_box(draw, *item, title_size=24, body_size=18)
        center_x = (item[0][0] + item[0][2]) // 2
        draw_arrow(draw, (900, 335), (center_x, item[0][1]), color=item[4], width=3)

    dependencies = [
        ((25, 660, 345, 780), "VectorRetriever", "Chroma PersistentClient", PALE_BLUE, BLUE),
        ((370, 660, 690, 780), "FinancialSentimentAnalyzer", "keyword sentiment + sample news", PALE_GOLD, GOLD),
        ((715, 660, 1035, 780), "RiskScorer", "wraps CentroidClassifier", PALE_RED, RED),
        ((1060, 660, 1380, 780), "PricePredictor", "wraps CentroidClassifier", PALE_GREEN, GREEN),
        ((1405, 660, 1775, 780), "AgentResult / EvidenceItem", "shared dataclass contract", LIGHT, MUTED),
    ]
    for item in dependencies:
        draw_box(draw, *item, title_size=23, body_size=18)
    for i in range(4):
        top = agents[i][0]
        bottom = dependencies[i][0]
        draw_arrow(draw, ((top[0] + top[2]) // 2, top[3]), ((bottom[0] + bottom[2]) // 2, bottom[1]), color=dependencies[i][4], width=3)
    for item in agents:
        top = item[0]
        draw.line([((top[0] + top[2]) // 2, top[3]), (1590, 660)], fill="#" + MUTED, width=2)

    draw_box(draw, (300, 900, 720, 1025), "ExplainabilityAgent", "Consumes all outputs and grounded evidence", PALE_TEAL, TEAL)
    draw_box(draw, (790, 900, 1180, 1025), "GenAIProvider", "abstract generation interface", PALE_BLUE, BLUE)
    draw_box(draw, (1250, 850, 1735, 945), "LocalGenAIProvider", "returns deterministic grounded prompt", LIGHT, MUTED)
    draw_box(draw, (1250, 990, 1735, 1085), "OpenAIGenAIProvider", "Responses API when key is configured", PALE_GREEN, GREEN)
    draw_arrow(draw, (1590, 780), (510, 900), color=TEAL, width=3)
    draw_arrow(draw, (720, 962), (790, 962), color=TEAL)
    draw_arrow(draw, (1180, 962), (1250, 897), color=MUTED)
    draw_arrow(draw, (1180, 962), (1250, 1037), color=GREEN)
    draw_label(draw, (900, 1145), "Composition arrows show runtime calls. All agents expose a stable AgentResult contract to reduce coupling.", 23, MUTED, True, "ma")
    return save_diagram(image, "05_class_flow.png")


def create_build_class_flow():
    image = Image.new("RGB", (1800, 1040), "white")
    draw = ImageDraw.Draw(image)
    draw_label(draw, (900, 40), "Data, Training, and RAG Build Classes", 42, NAVY, True, "ma")

    draw_box(draw, (70, 130, 470, 260), "AppSettings", "Environment-driven paths and providers", LIGHT, MUTED)
    draw_box(draw, (700, 130, 1100, 260), "DataStore", "Initialize and persist layered artifacts", PALE_BLUE, BLUE)
    draw_arrow(draw, (470, 195), (700, 195), color=MUTED)

    draw_box(draw, (70, 390, 470, 525), "US Equity Universe Builder", "Load, clean, deduplicate, cache", PALE_BLUE, BLUE)
    draw_box(draw, (700, 390, 1100, 525), "Phase1DataPipeline", "Generate features and labels", PALE_TEAL, TEAL)
    draw_box(draw, (1330, 390, 1730, 525), "Phase1ModelTrainer", "Split, train, evaluate, save", PALE_GOLD, GOLD)
    draw_arrow(draw, (270, 260), (270, 390), color=MUTED)
    draw_arrow(draw, (900, 260), (900, 390), color=BLUE)
    draw_arrow(draw, (470, 457), (700, 457), color=BLUE)
    draw_arrow(draw, (1100, 457), (1330, 457), color=TEAL)

    draw_box(draw, (1190, 650, 1450, 770), "RiskScorer", "Centroid model wrapper", PALE_RED, RED)
    draw_box(draw, (1490, 650, 1750, 770), "PricePredictor", "Centroid model wrapper", PALE_GREEN, GREEN)
    draw_arrow(draw, (1530, 525), (1320, 650), color=RED)
    draw_arrow(draw, (1530, 525), (1620, 650), color=GREEN)

    draw_box(draw, (70, 650, 390, 770), "DocumentLoader", "Collect, extract, chunk", PALE_BLUE, BLUE)
    draw_box(draw, (450, 650, 770, 770), "HashingEmbedder", "384-dimensional vectors", PALE_TEAL, TEAL)
    draw_box(draw, (830, 650, 1130, 770), "VectorRetriever", "Persist/query Chroma", PALE_GREEN, GREEN)
    draw_arrow(draw, (390, 710), (450, 710), color=BLUE)
    draw_arrow(draw, (770, 710), (830, 710), color=TEAL)

    draw_box(
        draw,
        (330, 850, 870, 980),
        "RAGSystem",
        "Index documents; retrieve evidence; synthesize answer",
        PALE_GOLD,
        GOLD,
        title_size=28,
        body_size=19,
    )
    draw_arrow(draw, (230, 770), (500, 850), color=BLUE)
    draw_arrow(draw, (980, 770), (700, 850), color=GREEN)
    draw_label(draw, (900, 1015), "Build-time classes produce the persistent assets consumed by the runtime agent graph.", 23, MUTED, True, "ma")
    return save_diagram(image, "06_build_class_flow.png")


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("FINANCIAL INTELLIGENCE ENGINE")
    set_run_font(r, size=11, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Application Architecture")
    set_run_font(r, size=30, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Phase 1 System Design, Class Flow, RAG, ML, and Deployment Path")
    set_run_font(r, size=14, color=BLUE)

    add_callout(
        doc,
        "Purpose",
        "A mentor-ready technical guide to the implemented Phase 1 application and the architecture that allows it to evolve into a cloud financial intelligence product.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    doc.add_paragraph()
    metadata = [
        ("Document status", "Phase 1 implementation baseline"),
        ("Prepared for", "Mentor and academic review"),
        ("Implementation date", "28 June 2026"),
        ("Application scope", "US equities, 10K-company universe"),
    ]
    add_table(doc, ["Field", "Value"], metadata, [2300, 7060], header_fill=LIGHT, font_size=10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run("Decision support only. The application does not provide regulated financial advice.")
    set_run_font(r, size=9.5, color=MUTED, italic=True)


def add_snapshot(doc):
    page_break(doc)
    add_heading(doc, "Executive Snapshot", 1)
    add_body(
        doc,
        "The Financial Intelligence Engine is a local-first, cloud-ready research and decision-support prototype. It combines document retrieval, persistent vector search, baseline machine-learning inference, specialist agents, transparent decision scoring, and an explainable final report.",
    )
    add_callout(
        doc,
        "Architecture in one sentence",
        "A browser question is routed through a Python API to an orchestrator, grounded with RAG evidence, enriched with sentiment and trained risk/forecast signals, scored by a decision agent, and rendered as an explainable report.",
    )

    add_heading(doc, "Current Evidence of Implementation", 2)
    metrics = [
        ("US equity universe", "10,000 tickers", "Cached Nasdaq Trader and SEC source files"),
        ("ML feature table", "10,000 rows x 20 columns", "Deterministic Phase 1 feature contract"),
        ("RAG vector index", "478 chunks", "Persistent ChromaDB collection"),
        ("Risk model", "80.05% accuracy; 0.8039 macro F1", "Low / Medium / High classifier"),
        ("Forecast model", "68.95% accuracy; 0.5914 macro F1", "Down / Flat / Up classifier"),
        ("Frontend", "5 views and 6 API routes", "Local dashboard at 127.0.0.1:8000"),
    ]
    add_table(doc, ["Capability", "Current result", "Implementation"], metrics, [2200, 2600, 4560], header_fill=PALE_BLUE, font_size=9)

    add_heading(doc, "Contents", 2)
    contents = [
        "1. System architecture overview",
        "2. Runtime query and agent flow",
        "3. Persistent RAG architecture",
        "4. Machine-learning training and inference",
        "5. Runtime and build-time class flow",
        "6. Data, storage, frontend, and API contracts",
        "7. Deployment evolution and Phase 2 path",
        "8. Current boundaries and mentor demo runbook",
    ]
    for item in contents:
        add_bullet(doc, item)


def add_system_overview(doc, system_image):
    page_break(doc)
    add_heading(doc, "1. System Architecture Overview", 1)
    add_body(
        doc,
        "The application follows a layered architecture. The browser and API form the interaction boundary; the orchestrator controls workflow; specialist agents produce evidence and predictions; the decision and explainability layers turn those signals into an auditable response.",
    )
    add_diagram(
        doc,
        system_image,
        "Phase 1 architecture from browser and API through orchestrated specialist agents, decision scoring, explainability, and final report.",
    )
    add_caption(doc, "Figure 1. Implemented Phase 1 end-to-end architecture.")

    add_heading(doc, "Key Design Decisions", 2)
    for text in [
        "Local-first operation: the complete demo works without a cloud API key.",
        "Persistent intelligence assets: vectors, features, models, metrics, and reports survive process restarts.",
        "Provider abstraction: GenAI can switch from deterministic local mode to OpenAI through environment configuration.",
        "Stable agent contract: every specialist returns AgentResult with status, data, evidence, confidence, and warnings.",
        "Transparent scoring: the decision layer exposes its input weights and recommendation thresholds.",
    ]:
        add_bullet(doc, text)


def add_layer_table(doc):
    page_break(doc)
    add_heading(doc, "2. Layered Component Architecture", 1)
    layers = [
        ("Presentation", "frontend/index.html, app.js, styles.css", "Dashboard, query form, universe, analysis, outputs"),
        ("HTTP/API", "DashboardHandler", "Routes GET endpoints, handles POST /api/ask, serves files"),
        ("Application", "OrchestratorAgent", "Selects tickers and executes one coordinated workflow"),
        ("Specialist agents", "Retriever, Sentiment, Risk, Forecast", "Produce evidence and structured analytic signals"),
        ("Decision", "DecisionAgent", "Combines signals into score and recommendation"),
        ("Explainability", "ExplainabilityAgent", "Produces grounded report, caveats, evidence, and disclaimer"),
        ("RAG", "DocumentLoader, HashingEmbedder, VectorRetriever, RAGSystem", "Indexing, persistent vectors, retrieval, answer synthesis"),
        ("ML", "Phase1ModelTrainer, RiskScorer, PricePredictor", "Training, evaluation, persistence, prediction"),
        ("Data/storage", "USEquityUniverseBuilder, Phase1DataPipeline, DataStore", "Universe, features, labels, layered artifacts"),
        ("Configuration", "AppSettings", "Environment, paths, provider, model, database URL"),
    ]
    add_table(doc, ["Layer", "Primary classes/files", "Responsibility"], layers, [1500, 3180, 4680], header_fill=PALE_TEAL, font_size=8.7)

    add_heading(doc, "Control Versus Data Flow", 2)
    add_body(
        doc,
        "Control flow is owned by OrchestratorAgent. Data flow is carried through AgentResult objects, EvidenceItem objects, Pandas DataFrames, persisted CSV/JSON files, Chroma documents and embeddings, and serialized PKL models.",
    )
    add_callout(
        doc,
        "Important implementation detail",
        "The conceptual diagram shows specialist agents as a group. In the current Python implementation, calls are sequential: retrieval -> sentiment -> risk -> forecast -> decision -> explanation. Forecast is logically independent and can be parallelized later.",
        fill=PALE_GOLD,
        accent=GOLD,
    )


def add_runtime_section(doc, runtime_image):
    page_break(doc)
    add_heading(doc, "3. Runtime Query and Agent Flow", 1)
    add_diagram(
        doc,
        runtime_image,
        "Runtime request sequence from POST api ask through retrieval, sentiment, risk, forecast, decision, explainability, and browser rendering.",
    )
    add_caption(doc, "Figure 2. Runtime request path from browser question to explainable response.")

    add_heading(doc, "Request Sequence", 2)
    steps = [
        "The browser submits a question to POST /api/ask.",
        "DashboardHandler validates the payload and extracts valid ticker symbols from the cached universe.",
        "OrchestratorAgent creates the specialist agents and begins one coordinated run.",
        "RetrieverAgent queries ChromaDB; if the index is absent or unavailable, it falls back to local hybrid retrieval.",
        "SentimentAgent combines generated news indicators with sentiment from retrieved evidence.",
        "RiskAgent and ForecastAgent load trained artifacts and select the requested ticker feature row.",
        "DecisionAgent applies transparent weights and maps the score to Strong Buy, Buy, Hold, Sell, or Strong Sell.",
        "ExplainabilityAgent creates a grounded prompt and delegates to the configured GenAIProvider.",
        "The API returns the report, suggestions, agent summaries, and educational disclaimer.",
    ]
    for step in steps:
        add_number(doc, step)


def add_rag_section(doc, rag_image):
    page_break(doc)
    add_heading(doc, "4. Persistent RAG Architecture", 1)
    add_diagram(
        doc,
        rag_image,
        "RAG indexing and query paths showing document loading, chunking, hashing embeddings, persistent ChromaDB, evidence retrieval, and answer generation.",
    )
    add_caption(doc, "Figure 3. Indexing and query paths for the Phase 1 RAG subsystem.")

    rag_details = [
        ("Supported input", "PDF, Markdown, text, CSV, and JSON"),
        ("Chunking", "900 characters with 120-character overlap"),
        ("Embedding", "Local deterministic token hashing; normalized 384-dimensional vectors"),
        ("Vector database", "ChromaDB PersistentClient"),
        ("Persistence path", "data/vectors/chroma/"),
        ("Collection", "phase1_financial_docs"),
        ("Retrieval output", "Text, source metadata, chunk ID, distance, and normalized relevance score"),
        ("Generation", "Local extractive fallback or OpenAI Responses API synthesis"),
    ]
    add_table(doc, ["RAG concern", "Current implementation"], rag_details, [2200, 7160], header_fill=PALE_BLUE, font_size=9.2)

    add_callout(
        doc,
        "Why RAG and ML both exist",
        "RAG supplies source-grounded context and explains why. ML supplies learned numerical risk and direction signals and estimates what is likely. The decision agent combines both forms of intelligence.",
        fill=PALE_TEAL,
        accent=TEAL,
    )


def add_ml_section(doc, ml_image):
    page_break(doc)
    add_heading(doc, "5. Machine-Learning Architecture", 1)
    add_diagram(
        doc,
        ml_image,
        "Machine-learning build and inference paths for the US equity universe, feature pipeline, model trainer, risk model, forecast model, and decision output.",
    )
    add_caption(doc, "Figure 4. Separate build-time training and runtime inference paths.")

    add_heading(doc, "Model Contracts", 2)
    models = [
        ("RiskScorer", "10 numeric features", "Low / Medium / High", "phase1_risk_model.pkl"),
        ("PricePredictor", "8 numeric features", "Down / Flat / Up", "phase1_forecast_model.pkl"),
        ("CentroidClassifier", "Standardized numeric matrix", "Nearest class centroid probabilities", "Embedded in both PKL files"),
    ]
    add_table(doc, ["Class", "Input", "Output", "Persistence"], models, [1800, 2400, 2780, 2380], header_fill=PALE_GOLD, font_size=8.8)

    add_heading(doc, "Training Method", 2)
    for item in [
        "Feature data is shuffled with random_state=42 and split 80% training / 20% testing.",
        "Features are standardized using training-set mean and standard deviation.",
        "Each class centroid is learned in standardized feature space.",
        "Prediction probabilities are derived from exponentiated negative distances.",
        "Accuracy, macro F1, weighted F1, and per-class precision/recall/F1 are persisted as JSON.",
    ]:
        add_bullet(doc, item)

    add_callout(
        doc,
        "Scientific boundary",
        "The current 10K feature table and labels are deterministic Phase 1 proxies designed to prove the pipeline at scale. They are not yet a backtested production market dataset. Phase 2 should replace proxies with point-in-time OHLCV, fundamentals, filings, and news while preventing look-ahead bias.",
        fill=PALE_RED,
        accent=RED,
    )


def add_class_flow_section(doc, class_image, build_class_image):
    page_break(doc)
    add_heading(doc, "6. Runtime Class Flow", 1)
    add_diagram(
        doc,
        class_image,
        "Runtime class dependency map for DashboardHandler, OrchestratorAgent, specialist agents, shared schemas, model wrappers, explainability, and GenAI providers.",
    )
    add_caption(doc, "Figure 5. Runtime class composition and dependency flow.")

    add_heading(doc, "Shared Result Contract", 2)
    add_body(
        doc,
        "EvidenceItem holds source, text, score, and metadata. AgentResult standardizes agent_name, status, summary, data, evidence, confidence, and warnings. This contract lets the orchestrator combine specialist outputs without depending on each agent's internal implementation.",
    )

    page_break(doc)
    add_heading(doc, "7. Data, Training, and RAG Build Classes", 1)
    add_diagram(
        doc,
        build_class_image,
        "Build-time class dependency map for configuration, storage, universe building, feature engineering, training, RAG indexing, and persisted model assets.",
    )
    add_caption(doc, "Figure 6. Build-time classes and the persistent assets they create.")

    class_catalog = [
        ("AppSettings", "config/settings.py", "Resolves local/cloud configuration"),
        ("DataStore", "storage/data_store.py", "Owns raw, processed, feature, vector, model, and report paths"),
        ("USEquityUniverseBuilder", "data_layer/company_universe.py", "Builds and caches the US ticker universe"),
        ("Phase1DataPipeline", "data_layer/phase1_data_pipeline.py", "Creates model-ready features and labels"),
        ("Phase1ModelTrainer", "ml_models/phase1_trainer.py", "Trains, evaluates, and saves both classifiers"),
        ("RAGSystem", "rag_layer/rag_system.py", "Provides standalone index and query API"),
        ("VectorRetriever", "rag_layer/retriever.py", "Owns Chroma collection and similarity operations"),
        ("GenAIProvider", "genai_layer/provider.py", "Abstracts local and OpenAI report generation"),
    ]
    add_table(doc, ["Class", "Module", "Role"], class_catalog, [2200, 3000, 4160], header_fill=PALE_TEAL, font_size=8.7)


def add_data_api_section(doc):
    page_break(doc)
    add_heading(doc, "8. Data, Storage, Frontend, and API Contracts", 1)
    add_heading(doc, "Storage Layout", 2)
    storage = [
        ("data/raw/", "Unchanged Nasdaq Trader / SEC source downloads"),
        ("data/processed/", "Cleaned 10K universe, summaries, and manifests"),
        ("data/features/", "ML-ready phase1_model_features.csv"),
        ("data/vectors/chroma/", "Persistent ChromaDB collection"),
        ("models/", "Risk and forecast PKL artifacts plus metrics JSON"),
        ("reports/", "Generated Markdown reports"),
        ("data/*.csv / *.json", "Dashboard analysis and demonstration outputs"),
    ]
    add_table(doc, ["Location", "Responsibility"], storage, [2600, 6760], header_fill=PALE_BLUE, font_size=9)

    add_heading(doc, "Frontend API", 2)
    endpoints = [
        ("GET", "/api/summary", "Counts, model metrics, recommendations, files"),
        ("GET", "/api/universe", "Search and list cached US equity universe"),
        ("GET", "/api/recommendations", "Rank investment report rows"),
        ("GET", "/api/sentiment", "Return sentiment leaders"),
        ("GET", "/api/news", "Return latest demo news rows"),
        ("POST", "/api/ask", "Run the complete multi-agent workflow"),
    ]
    add_table(doc, ["Method", "Route", "Purpose"], endpoints, [1100, 2800, 5460], header_fill=LIGHT, font_size=9)

    add_heading(doc, "Configuration Boundary", 2)
    for item in [
        "FIE_ENV selects local or cloud-oriented behavior.",
        "FIE_DATA_ROOT and related variables redirect storage without changing business logic.",
        "FIE_GENAI_PROVIDER selects local or openai.",
        "FIE_GENAI_MODEL selects the hosted model.",
        "OPENAI_API_KEY must be injected as a secret and must never be committed.",
        "FIE_DATABASE_URL is prepared for a future SQLite/PostgreSQL transition.",
    ]:
        add_bullet(doc, item)


def add_deployment_section(doc):
    page_break(doc)
    add_heading(doc, "9. Deployment Evolution", 1)
    add_heading(doc, "Current Phase 1 Deployment", 2)
    current = [
        ("Client", "Browser on the demonstration machine"),
        ("Application", "Python ThreadingHTTPServer at 127.0.0.1:8000"),
        ("Persistence", "Local folders, ChromaDB, CSV/JSON, PKL artifacts"),
        ("GenAI", "Local fallback; OpenAI optional through environment variables"),
        ("Operations", "Manual indexing, training, and server startup scripts"),
    ]
    add_table(doc, ["Concern", "Current implementation"], current, [2100, 7260], header_fill=PALE_BLUE, font_size=9.2)

    add_heading(doc, "Startup-Grade Target Architecture", 2)
    target = [
        ("Frontend", "React or Next.js behind CDN"),
        ("Backend", "FastAPI service in a Docker container"),
        ("Data", "PostgreSQL plus object storage for filings and market datasets"),
        ("Vector search", "Managed Qdrant, Pinecone, Weaviate, or cloud-persistent Chroma"),
        ("Model serving", "Versioned model registry and dedicated inference service"),
        ("Async work", "Queue and workers for SEC ingestion, indexing, training, and alerts"),
        ("Security", "Identity, authorization, secret manager, encryption, audit logging"),
        ("Observability", "Request traces, agent latency, retrieval quality, model drift"),
    ]
    add_table(doc, ["Layer", "Recommended evolution"], target, [1900, 7460], header_fill=PALE_TEAL, font_size=9)

    add_callout(
        doc,
        "Recommended progression",
        "First package the existing application in Docker with a persistent volume. Then separate the API and frontend, move structured data to PostgreSQL, move documents to object storage, introduce managed vector search, and finally add multi-tenant security and billing.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "Phase 2 Research Priorities", 2)
    for item in [
        "Point-in-time market and fundamental dataset with reproducible lineage.",
        "Real 10-K, 10-Q, earnings-call, and news ingestion with ticker/date/page metadata.",
        "Finance-specific embeddings, hybrid retrieval, reranking, and retrieval evaluation.",
        "Walk-forward model validation, calibration, drift monitoring, and ablation studies.",
        "Agent routing, parallel execution, confidence gating, and citation verification.",
        "Portfolio-level risk, personalization, alerts, and analyst workflow features.",
    ]:
        add_bullet(doc, item)


def add_status_runbook(doc):
    page_break(doc)
    add_heading(doc, "10. Current Boundaries and Demo Runbook", 1)
    add_heading(doc, "What Is Implemented Now", 2)
    status_rows = [
        ("Implemented", "10K US ticker universe and layered local storage"),
        ("Implemented", "10K x 20 Phase 1 feature and label table"),
        ("Implemented", "Trained risk and forecast baseline models with metrics"),
        ("Implemented", "Persistent ChromaDB with 384-dimensional local embeddings"),
        ("Implemented", "Retriever, Sentiment, Risk, Forecast, Decision, Explainability agents"),
        ("Implemented", "Provider-agnostic GenAI layer and OpenAI adapter"),
        ("Implemented", "Local frontend, REST-style endpoints, and generated reports"),
        ("Boundary", "Most feature values and news are deterministic demo proxies"),
        ("Boundary", "RAG corpus is project/research material, not yet a broad filing corpus"),
        ("Boundary", "HTTP server has no authentication, HTTPS, quotas, or tenant isolation"),
        ("Boundary", "The decision score is educational and not a trading recommendation"),
    ]
    add_table(doc, ["Status", "Capability or boundary"], status_rows, [1700, 7660], header_fill=LIGHT, font_size=9)

    add_heading(doc, "Mentor Demo Sequence", 2)
    commands = [
        ("1", "Initialize storage", "python setup_phase1_storage.py --min-count 10000"),
        ("2", "Build features", "python build_phase1_data.py --min-count 10000"),
        ("3", "Train models", "python train_phase1_models.py"),
        ("4", "Index RAG", "python index_phase1_rag.py"),
        ("5", "Test RAG", 'python query_phase1_rag.py "What is implemented in Phase 1?"'),
        ("6", "Run agent demo", "python demo_phase1_agents.py"),
        ("7", "Start frontend", "python frontend/server.py"),
        ("8", "Open dashboard", "http://127.0.0.1:8000"),
    ]
    add_table(doc, ["Step", "Action", "Command"], commands, [700, 1900, 6760], header_fill=PALE_BLUE, font_size=8.8)

    add_heading(doc, "Suggested Mentor Explanation", 2)
    add_callout(
        doc,
        "Thirty-second summary",
        "Phase 1 proves an end-to-end financial intelligence workflow at 10K-company scale. Persistent RAG grounds the answer in documents, trained baseline models provide risk and forecast signals, specialist agents keep responsibilities modular, and the explainability layer produces a traceable report. Phase 2 replaces proxy data with point-in-time financial data and evaluates retrieval, models, and agents rigorously.",
        fill=PALE_TEAL,
        accent=TEAL,
    )


def build():
    diagrams = {
        "system": create_system_architecture(),
        "runtime": create_runtime_flow(),
        "rag": create_rag_flow(),
        "ml": create_ml_flow(),
        "classes": create_class_flow(),
        "build_classes": create_build_class_flow(),
    }

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Financial Intelligence Engine - Application Architecture and Class Flow"
    doc.core_properties.subject = "Phase 1 technical architecture"
    doc.core_properties.author = "Financial Intelligence Engine Project"
    doc.core_properties.keywords = "RAG, machine learning, multi-agent, financial intelligence, architecture"

    add_cover(doc)
    add_snapshot(doc)
    add_system_overview(doc, diagrams["system"])
    add_layer_table(doc)
    add_runtime_section(doc, diagrams["runtime"])
    add_rag_section(doc, diagrams["rag"])
    add_ml_section(doc, diagrams["ml"])
    add_class_flow_section(doc, diagrams["classes"], diagrams["build_classes"])
    add_data_api_section(doc)
    add_deployment_section(doc)
    add_status_runbook(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
